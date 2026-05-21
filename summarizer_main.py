from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lex_rank import LexRankSummarizer
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from pdfminer.high_level import extract_text as extract_pdf_text
import os

# This global will hold the most recent summary for voice
last_summary_text = ""

def read_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        elif ext == ".docx":
            doc = Document(file_path)
            return "\n".join([p.text for p in doc.paragraphs])
        elif ext == ".pdf":
            return extract_pdf_text(file_path)
        else:
            raise ValueError("Unsupported file format.")
    except UnicodeDecodeError:
        raise UnicodeDecodeError("File encoding must be UTF-8.")
    except Exception as e:
        raise e

def summarize_text(text):
    parser = PlaintextParser.from_string(text, Tokenizer("english"))
    summarizer = LexRankSummarizer()

    total_sentences = len(list(parser.document.sentences))
    if total_sentences > 100:
        sentence_count = 25
    elif total_sentences > 50:
        sentence_count = 15
    else:
        sentence_count = 10

    summary = summarizer(parser.document, sentence_count)
    return "\n".join(str(sentence) for sentence in summary)

def export_to_docx(summary, path):
    doc = Document()
    doc.add_heading("Summary", 0)
    doc.add_paragraph(summary)
    doc.save(path)

def export_to_pdf(summary, path):
    pdf = canvas.Canvas(path, pagesize=letter)
    width, height = letter
    text_object = pdf.beginText(40, height - 40)
    text_object.setFont("Times-Roman", 11)

    for line in summary.splitlines():
        if text_object.getY() < 40:
            pdf.drawText(text_object)
            pdf.showPage()
            text_object = pdf.beginText(40, height - 40)
            text_object.setFont("Times-Roman", 11)
        text_object.textLine(line)

    pdf.drawText(text_object)
    pdf.save()

def summarize_and_export(input_path, export_format, export_path):
    global last_summary_text
    text = read_file(input_path)
    summary_text = summarize_text(text)
    last_summary_text = summary_text

    if export_format == "docx":
        export_to_docx(summary_text, export_path)
    elif export_format == "pdf":
        export_to_pdf(summary_text, export_path)
    else:
        with open(export_path, "w", encoding="utf-8") as f:
            f.write(summary_text)

def get_latest_summary_text():
    return last_summary_text
